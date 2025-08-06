import time
import ollama
import pyttsx3
import json
import pyaudio
import numpy as np
from vosk import Model, KaldiRecognizer, SetLogLevel
import tkinter as tk
from tkinter import scrolledtext, font, ttk, filedialog, messagebox
import threading
import queue
import wave
import os
import math
import fitz  # PyMuPDF for PDF parsing
import docx  # For DOCX parsing
import re
import pandas as pd
import random
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from datetime import datetime
import sys

# 设置UTF-8编码支持
if sys.platform.startswith('win'):
    import locale
    try:
        locale.setlocale(locale.LC_ALL, 'zh_CN.UTF-8')
    except:
        try:
            locale.setlocale(locale.LC_ALL, 'Chinese_China.936')
        except:
            pass

# 确保标准输出使用UTF-8编码
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')


# JD职位描述解析模块
class JDAnalyzer:
    """职位描述分析器"""
    def __init__(self):
        self.jd_data = {}
        
    def parse_jd(self, jd_text):
        """解析JD文本，提取关键信息"""
        self.jd_data = {
            "position": "",
            "requirements": [],
            "responsibilities": [],
            "skills": [],
            "experience": "",
            "education": "",
            "keywords": []
        }
        
        # 提取职位名称
        lines = jd_text.split('\n')
        if lines:
            self.jd_data["position"] = lines[0].strip()
        
        # 关键词匹配提取
        requirements_keywords = ["要求", "任职要求", "岗位要求", "职位要求"]
        responsibilities_keywords = ["职责", "工作职责", "岗位职责", "工作内容"]
        skills_keywords = ["技能", "技术要求", "专业技能", "掌握"]
        
        current_section = ""
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
                
            # 判断当前段落类型
            if any(keyword in line for keyword in requirements_keywords):
                current_section = "requirements"
                continue
            elif any(keyword in line for keyword in responsibilities_keywords):
                current_section = "responsibilities"
                continue
            elif any(keyword in line for keyword in skills_keywords):
                current_section = "skills"
                continue
            
            # 添加到对应段落
            if current_section == "requirements":
                self.jd_data["requirements"].append(line)
            elif current_section == "responsibilities":
                self.jd_data["responsibilities"].append(line)
            elif current_section == "skills":
                self.jd_data["skills"].append(line)
        
        # 提取技术关键词
        tech_keywords = ["Java", "Python", "JavaScript", "React", "Vue", "Node.js", "Spring", "MySQL", 
                        "Redis", "Docker", "Kubernetes", "微服务", "分布式", "高并发", "架构设计",
                        "算法", "数据结构", "设计模式", "Linux", "Git", "Jenkins", "DevOps"]
        
        jd_lower = jd_text.lower()
        for keyword in tech_keywords:
            if keyword.lower() in jd_lower:
                self.jd_data["keywords"].append(keyword)
        
        return self.jd_data
    
    def get_jd_summary(self):
        """获取JD摘要"""
        if not self.jd_data:
            return "未解析JD信息"
        
        return f"""
职位: {self.jd_data['position']}
技术要求: {', '.join(self.jd_data['keywords'][:8])}
要求数量: {len(self.jd_data['requirements'])}项
职责数量: {len(self.jd_data['responsibilities'])}项
"""


# 三阶段面试管理器
class ThreeStageInterviewManager:
    """三阶段面试管理器：非技术问题、经历类问题、技术类问题"""
    def __init__(self):
        self.stages = ["非技术问题", "经历类问题", "技术类问题"]
        self.current_stage = 0
        self.stage_questions = {
            "非技术问题": [],
            "经历类问题": [],
            "技术类问题": []
        }
        self.stage_scores = {
            "非技术问题": [],
            "经历类问题": [],
            "技术类问题": []
        }
        self.questions_per_stage = [2, 3, 3]  # 每阶段默认问题数
        self.current_stage_question_count = 0
        
    def get_current_stage(self):
        """获取当前阶段"""
        if self.current_stage < len(self.stages):
            return self.stages[self.current_stage]
        return "面试结束"
    
    def should_advance_stage(self):
        """判断是否应该进入下一阶段"""
        return self.current_stage_question_count >= self.questions_per_stage[self.current_stage]
    
    def advance_to_next_stage(self):
        """进入下一阶段"""
        if self.current_stage < len(self.stages) - 1:
            self.current_stage += 1
            self.current_stage_question_count = 0
            return True
        return False
    
    def add_question_to_stage(self, question, stage=None):
        """添加问题到当前阶段"""
        if stage is None:
            stage = self.get_current_stage()
        if stage in self.stage_questions:
            self.stage_questions[stage].append(question)
            self.current_stage_question_count += 1
    
    def add_score_to_stage(self, score, stage=None):
        """添加评分到当前阶段"""
        if stage is None:
            stage = self.get_current_stage()
        if stage in self.stage_scores:
            self.stage_scores[stage].append(score)
    
    def get_stage_prompt(self, jd_data=None, resume_data=None):
        """获取当前阶段的提示"""
        current_stage = self.get_current_stage()
        
        if current_stage == "非技术问题":
            return self._get_non_technical_prompt(jd_data, resume_data)
        elif current_stage == "经历类问题":
            return self._get_experience_prompt(resume_data)
        elif current_stage == "技术类问题":
            return self._get_technical_prompt(jd_data)
        else:
            return "面试已结束"
    
    def _get_non_technical_prompt(self, jd_data, resume_data):
        """非技术问题阶段提示"""
        prompt = f"""
### 当前阶段：非技术问题（第{self.current_stage_question_count + 1}题）

这是面试的第一阶段，主要考察候选人的基本情况、沟通能力和职业规划。

**问题类型要求：**
1. 自我介绍类问题
2. 职业发展规划
3. 对公司和岗位的了解
4. 基本的工作态度和价值观

**基于简历信息生成问题：**
"""
        if resume_data:
            prompt += f"- 候选人姓名: {resume_data.get('name', '未知')}\n"
            prompt += f"- 工作经验: {len(resume_data.get('experience', []))}项\n"
            prompt += f"- 主要技能: {', '.join(resume_data.get('skills', [])[:3])}\n"
        
        if jd_data:
            prompt += f"- 目标职位: {jd_data.get('position', '未知')}\n"
        
        return prompt
    
    def _get_experience_prompt(self, resume_data):
        """经历类问题阶段提示"""
        prompt = f"""
### 当前阶段：经历类问题（第{self.current_stage_question_count + 1}题）

这是面试的第二阶段，深入挖掘候选人的项目经历和工作经验。

**问题类型要求：**
1. 深入挖掘简历中的项目经历
2. 具体技术实现细节
3. 遇到的挑战和解决方案
4. 团队协作和角色分工

**基于简历经历生成问题：**
"""
        if resume_data:
            if resume_data.get('projects'):
                prompt += "**项目经历：**\n"
                for project in resume_data['projects'][:3]:
                    prompt += f"- {project}\n"
            
            if resume_data.get('experience'):
                prompt += "**工作经验：**\n"
                for exp in resume_data['experience'][:3]:
                    prompt += f"- {exp}\n"
        
        return prompt
    
    def _get_technical_prompt(self, jd_data):
        """技术类问题阶段提示"""
        prompt = f"""
### 当前阶段：技术类问题（第{self.current_stage_question_count + 1}题）

这是面试的第三阶段，重点考察候选人的专业技术能力。

**问题类型要求：**
1. 基于JD要求的核心技术能力
2. 算法和数据结构（如果相关）
3. 系统设计和架构能力
4. 技术深度和广度

**基于JD要求生成问题：**
"""
        if jd_data:
            if jd_data.get('keywords'):
                prompt += f"**关键技术要求：** {', '.join(jd_data['keywords'])}\n"
            if jd_data.get('requirements'):
                prompt += "**技术要求细节：**\n"
                for req in jd_data['requirements'][:3]:
                    prompt += f"- {req}\n"
        
        return prompt
    
    def get_stage_summary(self):
        """获取各阶段总结"""
        summary = "### 三阶段面试总结：\n"
        for i, stage in enumerate(self.stages):
            scores = self.stage_scores[stage]
            if scores:
                avg_score = sum(scores) / len(scores)
                summary += f"**{stage}**: {len(scores)}题, 平均分: {avg_score:.2f}\n"
            else:
                summary += f"**{stage}**: 未完成\n"
        return summary
    
    def adjust_stage_questions(self, recent_scores):
        """根据表现动态调整后续阶段的问题数量"""
        if len(recent_scores) < 2:
            return
        
        avg_recent = sum(recent_scores[-2:]) / len(recent_scores[-2:])
        
        # 根据表现调整后续阶段问题数量（适应新的评分体系）
        if avg_recent >= 0.75:
            # 表现优秀，可以适当增加技术问题数量
            if self.current_stage < 2:
                self.questions_per_stage[2] = min(5, self.questions_per_stage[2] + 1)
        elif avg_recent < 0.5:
            # 表现需要改善，适当减少技术问题，增加基础问题
            if self.current_stage < 2:
                self.questions_per_stage[2] = max(2, self.questions_per_stage[2] - 1)
                self.questions_per_stage[1] = min(4, self.questions_per_stage[1] + 1)


# 面试复盘和报告生成器
class InterviewReviewManager:
    """面试复盘和报告管理器"""
    def __init__(self):
        self.interview_records = []
        self.detailed_feedback = []
        self.overall_assessment = {}
        
    def add_qa_record(self, question, answer, score, feedback, stage):
        """添加问答记录"""
        record = {
            "timestamp": datetime.now().strftime("%H:%M:%S"),
            "stage": stage,
            "question": question,
            "answer": answer,
            "score": score,
            "feedback": feedback
        }
        self.interview_records.append(record)
    
    def generate_detailed_feedback(self, answer, question_context, score):
        """生成详细反馈"""
        try:
            feedback_prompt = f"""
请作为专业面试官，对以下回答给出详细的评价和改进建议：

问题背景：{question_context}
候选人回答：{answer}
当前评分：{score:.2f}

请从以下维度给出反馈：
1. 回答的完整性和准确性
2. 技术深度和理解程度
3. 表达清晰度和逻辑性
4. 具体的改进建议

请用简洁的中文回复，不超过100字。
"""
            
            response = ollama.chat(
                model="Jerrypoi/deepseek-r1-with-tool-calls:latest",
                messages=[{"role": "user", "content": feedback_prompt}]
            )
            
            return response['message']['content'].strip()
        except Exception as e:
            return f"技术掌握程度：{score:.1f}/1.0，建议加强相关技术的深入学习和实践。"
    
    def generate_overall_assessment(self, stage_manager, score_manager):
        """生成综合评估"""
        self.overall_assessment = {
            "interview_date": datetime.now().strftime("%Y年%m月%d日"),
            "total_questions": len(self.interview_records),
            "overall_score": sum([r["score"] for r in self.interview_records]) / len(self.interview_records) if self.interview_records else 0,
            "stage_performance": {},
            "strengths": [],
            "weaknesses": [],
            "improvement_suggestions": []
        }
        
        # 各阶段表现
        for stage in stage_manager.stages:
            stage_records = [r for r in self.interview_records if r["stage"] == stage]
            if stage_records:
                avg_score = sum([r["score"] for r in stage_records]) / len(stage_records)
                self.overall_assessment["stage_performance"][stage] = {
                    "questions": len(stage_records),
                    "avg_score": avg_score
                }
        
        # 生成优势、不足和建议
        self._analyze_performance()
        
        return self.overall_assessment
    
    def _analyze_performance(self):
        """分析面试表现"""
        if not self.interview_records:
            return
        
        # 分析各阶段表现
        stage_scores = {}
        for record in self.interview_records:
            stage = record["stage"]
            if stage not in stage_scores:
                stage_scores[stage] = []
            stage_scores[stage].append(record["score"])
        
        # 确定优势和不足（适应新的评分体系）
        for stage, scores in stage_scores.items():
            avg_score = sum(scores) / len(scores)
            if avg_score >= 0.7:
                self.overall_assessment["strengths"].append(f"{stage}表现优秀")
            elif avg_score < 0.55:
                self.overall_assessment["weaknesses"].append(f"{stage}需要加强")
        
        # 生成改进建议（适应新的评分体系）
        overall_score = self.overall_assessment["overall_score"]
        if overall_score < 0.55:
            self.overall_assessment["improvement_suggestions"].extend([
                "建议加强基础技术知识的学习",
                "多做项目实践，积累实际经验",
                "提高技术表达和沟通能力"
            ])
        elif overall_score < 0.75:
            self.overall_assessment["improvement_suggestions"].extend([
                "继续深化技术理解",
                "关注行业前沿技术发展",
                "提升系统设计和架构能力"
            ])
        else:
            self.overall_assessment["improvement_suggestions"].extend([
                "表现优秀，继续保持",
                "可以尝试挑战更高难度的技术领域",
                "分享经验，帮助他人成长"
            ])
    
    def export_to_pdf(self, filename, candidate_name, track_name):
        """导出面试记录为PDF"""
        try:
            # 注册中文字体
            if sys.platform.startswith('darwin'):  # macOS
                try:
                    pdfmetrics.registerFont(TTFont('PingFang', '/System/Library/Fonts/PingFang.ttc'))
                    chinese_font = 'PingFang'
                except:
                    try:
                        pdfmetrics.registerFont(TTFont('STHeiti', '/System/Library/Fonts/STHeiti Light.ttc'))
                        chinese_font = 'STHeiti'
                    except:
                        chinese_font = 'Helvetica'
            elif sys.platform.startswith('win'):  # Windows
                try:
                    pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                    chinese_font = 'SimSun'
                except:
                    try:
                        pdfmetrics.registerFont(TTFont('Microsoft YaHei', 'C:/Windows/Fonts/msyh.ttc'))
                        chinese_font = 'Microsoft YaHei'
                    except:
                        chinese_font = 'Helvetica'
            else:  # Linux
                try:
                    pdfmetrics.registerFont(TTFont('WenQuanYi', '/usr/share/fonts/wenquanyi/wqy-zenhei.ttc'))
                    chinese_font = 'WenQuanYi'
                except:
                    chinese_font = 'Helvetica'
            
            doc = SimpleDocTemplate(filename, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # 创建自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=chinese_font,
                fontSize=18,
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=chinese_font,
                fontSize=14,
                spaceAfter=12
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=chinese_font,
                fontSize=10,
                spaceAfter=6
            )
            
            content = []
            
            # 标题
            content.append(Paragraph(f"AI面试报告 - {candidate_name}", title_style))
            content.append(Spacer(1, 20))
            
            # 基本信息
            content.append(Paragraph("基本信息", heading_style))
            basic_info = [
                ["面试日期", self.overall_assessment.get("interview_date", "")],
                ["面试赛道", track_name],
                ["总题数", str(self.overall_assessment.get("total_questions", 0))],
                ["综合得分", f"{self.overall_assessment.get('overall_score', 0):.2f}/1.00"]
            ]
            
            basic_table = Table(basic_info, colWidths=[2*inch, 3*inch])
            basic_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(basic_table)
            content.append(Spacer(1, 20))
            
            # 完整评估报告
            content.append(Paragraph("评估报告", heading_style))
            if hasattr(self, 'full_evaluation'):
                content.append(Paragraph(self.full_evaluation, normal_style))
            content.append(Spacer(1, 20))
            
            # 各阶段表现
            content.append(Paragraph("各阶段表现", heading_style))
            stage_data = [["阶段", "题数", "平均分"]]
            for stage, performance in self.overall_assessment.get("stage_performance", {}).items():
                stage_data.append([
                    stage,
                    str(performance["questions"]),
                    f"{performance['avg_score']:.2f}"
                ])
            
            stage_table = Table(stage_data, colWidths=[2*inch, 1*inch, 1*inch])
            stage_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(stage_table)
            content.append(Spacer(1, 20))
            
            # 详细问答记录
            content.append(Paragraph("详细问答记录", heading_style))
            for i, record in enumerate(self.interview_records):
                content.append(Paragraph(f"第{i+1}题 [{record['stage']}]", normal_style))
                content.append(Paragraph(f"问题：{record['question']}", normal_style))
                content.append(Paragraph(f"回答：{record['answer'][:200]}...", normal_style))
                content.append(Paragraph(f"得分：{record['score']:.2f}", normal_style))
                content.append(Paragraph(f"反馈：{record['feedback']}", normal_style))
                content.append(Spacer(1, 10))
            
            # 综合评价
            content.append(Paragraph("综合评价", heading_style))
            
            strengths = self.overall_assessment.get("strengths", [])
            if strengths:
                content.append(Paragraph("优势：", normal_style))
                for strength in strengths:
                    content.append(Paragraph(f"• {strength}", normal_style))
            
            weaknesses = self.overall_assessment.get("weaknesses", [])
            if weaknesses:
                content.append(Paragraph("待改进：", normal_style))
                for weakness in weaknesses:
                    content.append(Paragraph(f"• {weakness}", normal_style))
            
            suggestions = self.overall_assessment.get("improvement_suggestions", [])
            if suggestions:
                content.append(Paragraph("改进建议：", normal_style))
                for suggestion in suggestions:
                    content.append(Paragraph(f"• {suggestion}", normal_style))
            
            doc.build(content)
            return True
            
        except Exception as e:
            print(f"PDF导出失败: {e}")
            return False


# 题库管理模块
class QuestionBankManager:
    """题库管理器，支持多赛道面试题库"""
    def __init__(self):
        self.question_banks = {}
        self.tracks = ["后端", "前端", "算法", "测试", "产品", "运营"]
        self.current_track = None
        self.load_question_banks()
    
    def load_question_banks(self):
        """加载所有赛道的题库数据"""
        for track in self.tracks:
            try:
                file_path = f"data/data_divided/{track}.csv"
                df = pd.read_csv(file_path, encoding='utf-8')
                
                # 处理不同的CSV格式
                questions = []
                if track in ["后端"]:
                    # 后端格式: 分类,问题
                    for _, row in df.iterrows():
                        if len(row) >= 2 and pd.notna(row.iloc[1]):
                            questions.append({
                                "category": str(row.iloc[0]).strip(),
                                "question": str(row.iloc[1]).strip(),
                                "company": ""
                            })
                elif track in ["前端", "运营"]:
                    # 前端/运营格式: 分类,问题,公司
                    for _, row in df.iterrows():
                        if len(row) >= 2 and pd.notna(row.iloc[1]) and "Unnamed" not in str(row.iloc[1]):
                            questions.append({
                                "category": str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else "",
                                "question": str(row.iloc[1]).strip(),
                                "company": str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ""
                            })
                else:
                    # 其他格式: 分类,问题 或 分类,问题,公司
                    for _, row in df.iterrows():
                        if len(row) >= 2 and pd.notna(row.iloc[1]):
                            questions.append({
                                "category": str(row.iloc[0]).strip(),
                                "question": str(row.iloc[1]).strip(),
                                "company": str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ""
                            })
                
                # 按难度分类题目
                self.question_banks[track] = self._categorize_by_difficulty(questions)
                print(f"已加载 {track} 题库: {len(questions)} 题")
                
            except Exception as e:
                print(f"加载 {track} 题库失败: {e}")
                self.question_banks[track] = {"简单": [], "中等": [], "困难": []}
    
    def _categorize_by_difficulty(self, questions):
        """根据问题内容和类别自动分类难度"""
        categorized = {"简单": [], "中等": [], "困难": []}
        
        # 定义难度关键词
        easy_keywords = ["自我介绍", "介绍", "背景", "经历", "基础", "了解", "是什么", "简单"]
        hard_keywords = ["架构", "设计", "优化", "性能", "复杂", "系统", "深入", "原理", "底层", "高级"]
        
        for q in questions:
            question_text = q["question"].lower()
            category_text = q["category"].lower()
            combined_text = question_text + " " + category_text
            
            # 判断难度
            if any(keyword in combined_text for keyword in easy_keywords):
                categorized["简单"].append(q)
            elif any(keyword in combined_text for keyword in hard_keywords):
                categorized["困难"].append(q)
            else:
                categorized["中等"].append(q)
        
        return categorized
    
    def set_track(self, track):
        """设置当前面试赛道"""
        if track in self.tracks:
            self.current_track = track
            return True
        return False
    
    def get_reference_questions(self, difficulty="中等", num_questions=3):
        """获取参考问题用于AI生成"""
        if not self.current_track or self.current_track not in self.question_banks:
            return []
        
        track_questions = self.question_banks[self.current_track]
        difficulty_questions = track_questions.get(difficulty, [])
        
        if not difficulty_questions:
            # 如果当前难度没有题目，从其他难度选择
            all_questions = []
            for diff_level in track_questions.values():
                all_questions.extend(diff_level)
            if all_questions:
                return random.sample(all_questions, min(num_questions, len(all_questions)))
            return []
        
        return random.sample(difficulty_questions, min(num_questions, len(difficulty_questions)))
    
    def get_track_summary(self):
        """获取当前赛道题库摘要"""
        if not self.current_track:
            return "未选择赛道"
        
        track_data = self.question_banks.get(self.current_track, {})
        summary = f"{self.current_track}赛道题库:\n"
        for difficulty, questions in track_data.items():
            summary += f"- {difficulty}: {len(questions)}题\n"
        
        return summary


# 评分系统和动态难度调整模块
class ScoreAndDifficultyManager:
    """评分系统和动态难度调整管理器"""
    def __init__(self):
        self.current_difficulty = "中等"  # 初始难度为中等
        self.score_history = []  # 历史评分记录
        self.difficulty_history = []  # 难度调整历史
        self.question_count = 0  # 问题计数
        
        # 难度级别定义
        self.difficulty_levels = {
            "简单": {
                "level": 1,
                "keywords": ["基础", "简单", "基本", "介绍"],
                "description": "基础概念和简单问题"
            },
            "中等": {
                "level": 2,
                "keywords": ["实际", "应用", "经验", "项目"],
                "description": "实际应用和项目经验"
            },
            "困难": {
                "level": 3,
                "keywords": ["深入", "复杂", "高级", "架构", "优化"],
                "description": "深度技术和复杂场景"
            }
        }
    
    def evaluate_response(self, user_response, question_context=""):
        """
        使用AI模型对用户回答进行评分
        返回0-1之间的分数
        """
        try:
            # 构建评分提示
            scoring_prompt = f"""
你是一个友善而专业的面试官，正在对候选人的回答进行评分。请根据面试的实际情况，给出公平合理的评分。

面试问题: {question_context}
候选人回答: {user_response}

评分指导原则：
- 如果候选人的回答合理、相关且表达清楚，应给予0.6-0.8的分数
- 如果回答特别出色、有深度或有创新见解，可给予0.8-1.0的分数  
- 如果回答基本合理但略显简单，可给予0.4-0.6的分数
- 只有在回答完全不相关、错误或无法理解时，才给予0.4以下的分数

请模拟真实面试场景，用人性化的角度来评判。大多数正常的回答都应该在0.5-0.8这个合理区间内。

请只返回一个0到1之间的小数作为评分，保留两位小数，不要包含任何其他文字。
"""
            
            # 调用模型进行评分
            response = ollama.chat(
                model="Jerrypoi/deepseek-r1-with-tool-calls:latest",
                messages=[{"role": "user", "content": scoring_prompt}]
            )
            
            score_text = response['message']['content'].strip()
            
            # 提取数字分数
            import re
            score_match = re.search(r'(\d*\.?\d+)', score_text)
            if score_match:
                score = float(score_match.group(1))
                # 确保分数在0-1范围内
                score = max(0.0, min(1.0, score))
            else:
                # 如果无法解析，默认给中等分数
                score = 0.5
            
            # 记录评分历史
            self.score_history.append(score)
            self.question_count += 1
            
            return score
            
        except Exception as e:
            print(f"评分过程出错: {e}")
            # 出错时返回中等分数
            return 0.5
    
    def adjust_difficulty(self, score):
        """
        根据评分调整问题难度
        """
        previous_difficulty = self.current_difficulty
        
        # 根据评分调整难度（适应更人性化的评分体系）
        if score >= 0.75:
            # 表现优秀，增加难度
            if self.current_difficulty == "简单":
                self.current_difficulty = "中等"
            elif self.current_difficulty == "中等":
                self.current_difficulty = "困难"
            # 已经是困难级别，保持不变
        elif score < 0.5:
            # 表现需要改善，降低难度
            if self.current_difficulty == "困难":
                self.current_difficulty = "中等"
            elif self.current_difficulty == "中等":
                self.current_difficulty = "简单"
            # 已经是简单级别，保持不变
        # 0.5 <= score < 0.75，保持当前难度
        
        # 记录难度调整历史
        self.difficulty_history.append({
            "question_num": self.question_count,
            "score": score,
            "previous_difficulty": previous_difficulty,
            "new_difficulty": self.current_difficulty
        })
        
        return self.current_difficulty
    
    def get_difficulty_prompt(self):
        """
        获取当前难度对应的提示语
        """
        current_level = self.difficulty_levels[self.current_difficulty]
        
        prompt = f"""
当前问题难度级别: {self.current_difficulty} (级别 {current_level['level']})
难度描述: {current_level['description']}
建议关键词: {', '.join(current_level['keywords'])}

请根据此难度级别生成相应的面试问题。

难度级别说明：
- 简单：基础概念、基本技能、入门级问题
- 中等：实际应用、项目经验、中级技术问题
- 困难：深度技术、复杂场景、高级架构问题
"""
        return prompt
    
    def get_score_summary(self):
        """
        获取评分摘要
        """
        if not self.score_history:
            return "暂无评分记录"
        
        avg_score = sum(self.score_history) / len(self.score_history)
        max_score = max(self.score_history)
        min_score = min(self.score_history)
        
        summary = f"""
评分摘要：
- 总问题数: {len(self.score_history)}
- 平均分: {avg_score:.2f}
- 最高分: {max_score:.2f}
- 最低分: {min_score:.2f}
- 当前难度: {self.current_difficulty}
"""
        return summary
    
    def get_difficulty_progression(self):
        """
        获取难度变化轨迹
        """
        if not self.difficulty_history:
            return "暂无难度调整记录"
        
        progression = "难度调整轨迹：\n"
        for record in self.difficulty_history:
            progression += f"问题{record['question_num']}: 得分{record['score']:.2f} -> {record['previous_difficulty']} → {record['new_difficulty']}\n"
        
        return progression



class DynamicPromptAdjuster:
    """动态提示调整模块，包含Triplet Filter和Demo Selector"""
    def __init__(self, resume_data):
        self.resume_data = resume_data
        self.key_entities = []  # 关键实体列表
        self.historical_entities = []  # 历史实体
        self.historical_acts = []  # 历史行为
        self.current_entities = []  # 当前实体
        self.current_acts = []  # 当前行为
        self.retained_triplets = []  # 保留的三元组
        self.max_triplets = 5  # 最大保留三元组数
    
    def extract_entities(self):
        """从简历中提取关键实体"""
        entities = []
        # 提取技能实体
        if self.resume_data.get("skills"):
            entities.extend(self.resume_data["skills"])
        
        # 提取项目经验中的关键技术
        if self.resume_data.get("projects"):
            for project in self.resume_data["projects"]:
                if "(" in project and ")" in project:
                    techs = project.split("(")[1].split(")")[0].split(",")
                    entities.extend([tech.strip() for tech in techs if len(tech.strip()) > 3])
        
        # 去重
        self.key_entities = list(set(entities))
        return self.key_entities
    
    def update_conversation_context(self, user_input, model_output):
        """更新对话上下文，提取实体和行为"""
        # 从用户输入中提取实体
        user_entities = self._extract_entities_from_text(user_input)
        self.historical_entities.extend(user_entities)
        
        # 从模型输出中提取行为（问题类型）
        model_act = self._classify_question_type(model_output)
        self.historical_acts.append(model_act)
        
        # 预测当前实体
        self.current_entities = self._predict_entities()
    
    def _extract_entities_from_text(self, text):
        """从文本中提取实体"""
        # 简单的实体匹配（实际应用中可以使用NER模型）
        extracted = []
        for entity in self.key_entities:
            if entity.lower() in text.lower():
                extracted.append(entity)
        return extracted
    
    def _classify_question_type(self, text):
        """分类问题类型（行为）"""
        question_types = {
            "技术": ["技术", "技能", "编程", "框架", "语言"],
            "项目": ["项目", "经验", "案例", "实施"],
            "行为": ["行为", "情景", "处理", "挑战"],
            "动机": ["动机", "为什么", "原因", "兴趣"],
            "基础": ["介绍", "背景", "教育", "经历"]
        }
        
        for act, keywords in question_types.items():
            for keyword in keywords:
                if keyword in text:
                    return act
        return "其他"
    
    def _predict_entities(self):
        """预测当前可能相关的实体"""
        # 简单的预测：使用最近提到的实体
        if self.historical_entities:
            return list(set(self.historical_entities[-3:]))
        return self.key_entities[:3]
    
    def triplet_filter(self):
        """三元组过滤器实现"""
        # 从实体生成初始三元组
        initial_triplets = self._generate_initial_triplets()
        
        # 计算实体频率
        entity_freq = {}
        for head, rel, tail in initial_triplets:
            entity_freq[head] = entity_freq.get(head, 0) + 1
            entity_freq[tail] = entity_freq.get(tail, 0) + 1
        
        # 迭代过滤过程
        retained_triplets = []
        tau = 1  # 初始阈值
        
        while True:
            # 应用过滤规则
            filtered = []
            for head, rel, tail in initial_triplets:
                min_freq = min(entity_freq.get(head, 0), entity_freq.get(tail, 0))
                if min_freq >= tau:
                    filtered.append((head, rel, tail))
            
            # 检查是否满足最大数量限制
            if len(filtered) <= self.max_triplets or tau > 5:  # tau上限防止无限循环
                retained_triplets = filtered
                break
            
            # 增加阈值
            tau += 1
        
        self.retained_triplets = retained_triplets
        return retained_triplets
    
    def _generate_initial_triplets(self):
        """生成初始三元组（简化版）"""
        triplets = []
        
        # 技能-项目关联
        if self.resume_data.get("skills") and self.resume_data.get("projects"):
            for skill in self.resume_data["skills"][:3]:
                for project in self.resume_data["projects"][:2]:
                    triplets.append((skill, "应用于", project))
        
        # 技能-经验关联
        if self.resume_data.get("skills") and self.resume_data.get("experience"):
            for skill in self.resume_data["skills"][:3]:
                for exp in self.resume_data["experience"][:2]:
                    triplets.append((skill, "用于", exp))
        
        # 项目-技能关联
        if self.resume_data.get("projects"):
            for project in self.resume_data["projects"][:3]:
                parts = project.split("(")
                if len(parts) > 1:
                    techs = parts[1].split(")")[0].split(",")
                    for tech in techs[:3]:
                        triplets.append((project, "使用技术", tech.strip()))
        
        return triplets
    
    def demo_selector(self):
        """演示选择器实现（简化版）"""
        # 在实际应用中，这里会有一个演示库和匹配算法
        # 这里返回一个固定的演示示例
        demo = """
        <面试示例>
        面试官: 请介绍一下你在XX项目中的角色和贡献。
        候选人: 在该项目中，我担任后端开发负责人，负责系统架构设计和核心模块实现。
        面试官: 你在项目中遇到的最大技术挑战是什么？如何解决的？
        候选人: 我们面临高并发场景下的性能问题，我通过引入Redis缓存和优化数据库查询解决了问题。
        """
        return demo

# 语音识别模块
class VoiceRecorder:
    def __init__(self, model):
        self.model = model
        self.rec = KaldiRecognizer(model, 16000)
        self.p = pyaudio.PyAudio()
        self.stream = None
        self.frames = []
        self.is_recording = False
        self.audio_format = pyaudio.paInt16
        self.channels = 1
        self.rate = 16000
        self.chunk = 1024
        self.start_time = 0

    def start_recording(self):
        if not self.is_recording:
            self.frames = []
            self.stream = self.p.open(
                format=self.audio_format,
                channels=self.channels,
                rate=self.rate,
                input=True,
                frames_per_buffer=self.chunk
            )
            self.is_recording = True
            self.start_time = time.time()
            print("Recording started...")
            threading.Thread(target=self._record).start()

    def _record(self):
        while self.is_recording:
            data = self.stream.read(self.chunk, exception_on_overflow=False)
            self.frames.append(data)

    def stop_recording(self):
        if self.is_recording:
            self.is_recording = False
            self.stream.stop_stream()
            self.stream.close()
            print("Recording stopped.")
            return self.process_recording()
        return None

    def process_recording(self):
        # 保存为临时WAV文件
        temp_file = "temp_recording.wav"
        wf = wave.open(temp_file, 'wb')
        wf.setnchannels(self.channels)
        wf.setsampwidth(self.p.get_sample_size(self.audio_format))
        wf.setframerate(self.rate)
        wf.writeframes(b''.join(self.frames))
        wf.close()

        # 使用Vosk进行识别
        result_text = self._recognize_speech(temp_file)
        # 删除临时文件
        os.remove(temp_file)
        return result_text

    def _recognize_speech(self, filename):
        wf = wave.open(filename, 'rb')
        if wf.getnchannels() != 1 or wf.getsampwidth() != 2 or wf.getcomptype() != "NONE":
            print("Audio file must be WAV format mono PCM.")
            return None
        
        recognizer = KaldiRecognizer(self.model, wf.getframerate())
        recognizer.SetWords(True)
        
        results = []
        while True:
            data = wf.readframes(4000)
            if len(data) == 0:
                break
            if recognizer.AcceptWaveform(data):
                result = json.loads(recognizer.Result())
                results.append(result.get("text", ""))
        
        final_result = json.loads(recognizer.FinalResult())
        results.append(final_result.get("text", ""))
        return ' '.join(results)

# 简历解析类
class ResumeParser:
    def __init__(self):
        self.resume_data = {
            "name": "",
            "contact": "",
            "education": [],
            "experience": [],
            "skills": [],
            "projects": [],
            "summary": ""
        }
    
    def parse_resume(self, file_path):
        """根据文件类型调用相应的解析方法"""
        if file_path.lower().endswith('.pdf'):
            return self.parse_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.parse_docx(file_path)
        else:
            return "不支持的文件格式，请上传PDF或DOCX文件"
    
    def parse_pdf(self, file_path):
        """解析PDF格式的简历"""
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            
            return self._extract_info(text)
        except Exception as e:
            return f"解析PDF失败: {str(e)}"
    
    def parse_docx(self, file_path):
        """解析DOCX格式的简历"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return self._extract_info(text)
        except Exception as e:
            return f"解析DOCX失败: {str(e)}"
    
    def _extract_info(self, text):
        """从文本中提取简历信息"""
        # 提取姓名
        name_match = re.search(r"^(.*?)\n", text)
        if name_match:
            self.resume_data["name"] = name_match.group(1).strip()
        
        # 提取联系方式
        phone_match = re.search(r"(\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4})", text)
        email_match = re.search(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})", text)
        self.resume_data["contact"] = f"电话: {phone_match.group(0) if phone_match else '未找到'} | 邮箱: {email_match.group(0) if email_match else '未找到'}"
        
        # 提取教育背景
        edu_section = re.search(r"(教育背景|教育经历|学历|Education)(.*?)(?=(工作经历|项目经历|技能|$))", text, re.DOTALL | re.IGNORECASE)
        if edu_section:
            edu_text = edu_section.group(2)
            self.resume_data["education"] = [line.strip() for line in edu_text.split('\n') if line.strip()]
        
        # 提取工作经历
        exp_section = re.search(r"(工作经历|工作经验|工作|Experience)(.*?)(?=(项目经历|技能|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if exp_section:
            exp_text = exp_section.group(2)
            self.resume_data["experience"] = [line.strip() for line in exp_text.split('\n') if line.strip()]
        
        # 提取技能
        skills_section = re.search(r"(技能|专业技能|技术能力|Skills)(.*?)(?=(项目经历|工作经历|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if skills_section:
            skills_text = skills_section.group(2)
            self.resume_data["skills"] = [line.strip() for line in skills_text.split('\n') if line.strip()]
        
        # 提取项目经历
        projects_section = re.search(r"(项目经历|项目经验|项目|Projects)(.*?)(?=(技能|工作经历|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if projects_section:
            projects_text = projects_section.group(2)
            self.resume_data["projects"] = [line.strip() for line in projects_text.split('\n') if line.strip()]
        
        # 生成摘要
        self.resume_data["summary"] = f"候选人: {self.resume_data['name']}，教育背景: {len(self.resume_data['education'])}项，工作经验: {len(self.resume_data['experience'])}项，技能: {len(self.resume_data['skills'])}项"
        
        return self.resume_data

    def extract_entities(self, file_path):
        """从文件中提取实体"""
        text = ""
        if file_path.lower().endswith('.pdf'):
            try:
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text()
            except Exception:
                return []
        elif file_path.lower().endswith('.docx'):
            try:
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception:
                return []
        else:
            return []
        
        # 提取技能实体
        entities = []
        skills_section = re.search(r"(技能|专业技能|技术能力|Skills)(.*?)(?=(项目经历|工作经历|教育背景|$))", 
                                  text, re.DOTALL | re.IGNORECASE)
        if skills_section:
            skills_text = skills_section.group(2)
            skills = [line.strip() for line in skills_text.split('\n') if line.strip()]
            entities.extend(skills)
        
        # 提取项目经验中的技术实体
        projects_section = re.search(r"(项目经历|项目经验|项目|Projects)(.*?)(?=(技能|工作经历|教育背景|$))", 
                                    text, re.DOTALL | re.IGNORECASE)
        if projects_section:
            projects_text = projects_section.group(2)
            projects = [line.strip() for line in projects_text.split('\n') if line.strip()]
            for project in projects:
                if "(" in project and ")" in project:
                    techs = project.split("(")[1].split(")")[0].split(",")
                    entities.extend([tech.strip() for tech in techs])
        
        return list(set(entities))


# Tkinter 窗口模块
class InteractiveTextApp:
    def __init__(self, root, solution):
        self.root = root
        self.solution = solution
        self.root.title("AI面试智能官")
        self.root.geometry("1000x700")
        self.root.configure(bg="#f0f0f0")
        self.dynamic_prompt_adjuster = None  # 动态提示调整器
        self.conversation_context = []  # 对话上下文        
        self.score_manager = ScoreAndDifficultyManager()  # 评分和难度管理器
        self.question_bank_manager = QuestionBankManager()  # 题库管理器
        self.stage_manager = ThreeStageInterviewManager()  # 三阶段面试管理器
        self.jd_analyzer = JDAnalyzer()  # JD分析器
        self.review_manager = InterviewReviewManager()  # 复盘管理器
        self.last_question = ""  # 保存最后一个问题用于评分
        self.selected_track = None  # 选择的面试赛道
        self.jd_data = None  # JD数据        
        # 创建主框架
        main_frame = tk.Frame(root, bg="#f0f0f0")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # 设置字体（优先使用中文字体）
        try:
            # 尝试使用中文字体
            if sys.platform.startswith('win'):
                self.default_family = "Microsoft YaHei"
            elif sys.platform.startswith('darwin'):  # macOS
                self.default_family = "PingFang SC"
            else:  # Linux
                self.default_family = "WenQuanYi Micro Hei"
        except:
            self.default_family = "Helvetica"
        
        self.custom_font = font.Font(family=self.default_family, size=16)
        self.small_font = font.Font(family=self.default_family, size=12)
        self.title_font = font.Font(family=self.default_family, size=18, weight="bold")

        # 创建标题
        title_label = tk.Label(
            main_frame,
            text="AI面试智能官",
            font=(self.default_family, 20, "bold"),
            bg="#f0f0f0",
            fg="#2c3e50"
        )
        title_label.pack(pady=10)

        # 创建顶部按钮框架
        button_frame = tk.Frame(main_frame, bg="#f0f0f0")
        button_frame.pack(fill=tk.X, pady=10)
        
        # 添加简历上传按钮
        self.upload_btn = tk.Button(
            button_frame,
            text="上传简历",
            command=self.upload_resume,
            font=self.small_font,
            bg="#3498db",
            fg="white",
            padx=10,
            pady=5,
            relief=tk.FLAT
        )
        self.upload_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加JD上传按钮
        self.upload_jd_btn = tk.Button(
            button_frame,
            text="上传JD",
            command=self.upload_jd,
            font=self.small_font,
            bg="#e67e22",
            fg="white",
            padx=10,
            pady=5,
            relief=tk.FLAT
        )
        self.upload_jd_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加赛道选择按钮
        self.track_select_btn = tk.Button(
            button_frame,
            text="选择赛道",
            command=self.select_track,
            font=self.small_font,
            bg="#9b59b6",
            fg="white",
            padx=10,
            pady=5,
            relief=tk.FLAT
        )
        self.track_select_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加面试开始按钮
        self.start_interview_btn = tk.Button(
            button_frame,
            text="开始面试",
            state=tk.DISABLED,  # 初始状态为禁用
            command=self.start_interview,
            font=self.small_font,
            bg="#2ecc71",
            fg="white",
            padx=10,
            pady=5,
            relief=tk.FLAT,
            state=tk.DISABLED  # 初始状态禁用
        )
        self.start_interview_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加面试结束按钮
        self.end_interview_btn = tk.Button(
            button_frame,
            text="结束面试",
            command=self.end_interview,
            font=self.small_font,
            bg="#e74c3c",
            fg="white",
            padx=10,
            pady=5,
            relief=tk.FLAT,
            state=tk.DISABLED
        )
        self.end_interview_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加面试复盘按钮
        self.review_btn = tk.Button(
            button_frame,
            text="面试复盘",
            command=self.show_interview_review,
            font=self.small_font,
            bg="#8e44ad",
            fg="white",
            padx=10,
            pady=5,
            relief=tk.FLAT,
            state=tk.DISABLED
        )
        self.review_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加导出PDF按钮
        self.export_pdf_btn = tk.Button(
            button_frame,
            text="导出PDF",
            command=self.export_interview_pdf,
            font=self.small_font,
            bg="#16a085",
            fg="white",
            padx=10,
            pady=5,
            relief=tk.FLAT,
            state=tk.DISABLED
        )
        self.export_pdf_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加候选人信息面板
        info_frame = tk.LabelFrame(
            main_frame,
            text="候选人信息",
            font=self.small_font,
            bg="#f0f0f0",
            bd=2,
            relief=tk.GROOVE
        )
        info_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.info_label = tk.Label(
            info_frame,
            text="请上传简历以查看候选人信息",
            font=self.small_font,
            bg="#f0f0f0",
            fg="#7f8c8d",
            justify=tk.LEFT,
            wraplength=900
        )
        self.info_label.pack(fill=tk.X, padx=10, pady=5)
        
        # 添加赛道选择状态面板
        track_frame = tk.LabelFrame(
            main_frame,
            text="面试赛道",
            font=self.small_font,
            bg="#f0f0f0",
            bd=2,
            relief=tk.GROOVE
        )
        track_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.track_info_label = tk.Label(
            track_frame,
            text="请选择面试赛道",
            font=("Helvetica", 12, "bold"),
            bg="#f0f0f0",
            fg="#9b59b6",
            justify=tk.LEFT,
            wraplength=900
        )
        self.track_info_label.pack(fill=tk.X, padx=10, pady=5)
        
        # 添加面试阶段状态面板
        stage_frame = tk.LabelFrame(
            main_frame,
            text="面试阶段",
            font=self.small_font,
            bg="#f0f0f0",
            bd=2,
            relief=tk.GROOVE
        )
        stage_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.stage_info_label = tk.Label(
            stage_frame,
            text="当前阶段: 等待开始",
            font=("Helvetica", 12, "bold"),
            bg="#f0f0f0",
            fg="#e67e22",
            justify=tk.LEFT
        )
        self.stage_info_label.pack(fill=tk.X, padx=10, pady=5)
        
        # 添加评分和难度状态面板
        score_frame = tk.LabelFrame(
            main_frame,
            text="面试状态",
            font=self.small_font,
            bg="#f0f0f0",
            bd=2,
            relief=tk.GROOVE
        )
        score_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 创建两列布局
        left_score_frame = tk.Frame(score_frame, bg="#f0f0f0")
        left_score_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        right_score_frame = tk.Frame(score_frame, bg="#f0f0f0")
        right_score_frame.pack(side=tk.RIGHT, fill=tk.X, expand=True, padx=5)
        
        # 当前难度显示（面试过程中不显示具体难度）
        self.difficulty_label = tk.Label(
            left_score_frame,
            text="系统智能调节中",
            font=(self.default_family, 12, "bold"),
            bg="#f0f0f0",
            fg="#3498db"
        )
        self.difficulty_label.pack(anchor=tk.W, pady=2)
        
        # 最新评分显示（隐藏，仅用于后台记录）
        self.latest_score_label = tk.Label(
            left_score_frame,
            text="面试进行中...",
            font=self.small_font,
            bg="#f0f0f0",
            fg="#2ecc71"
        )
        self.latest_score_label.pack(anchor=tk.W, pady=2)
        
        # 平均分显示（隐藏，仅用于后台记录）
        self.avg_score_label = tk.Label(
            right_score_frame,
            text="实时分析中...",
            font=self.small_font,
            bg="#f0f0f0",
            fg="#f39c12"
        )
        self.avg_score_label.pack(anchor=tk.W, pady=2)
        
        # 问题计数显示
        self.question_count_label = tk.Label(
            right_score_frame,
            text="问题数: 0",
            font=self.small_font,
            bg="#f0f0f0",
            fg="#9b59b6"
        )
        self.question_count_label.pack(anchor=tk.W, pady=2)

        # 创建可滚动的文本框
        text_frame = tk.Frame(main_frame, bg="#ffffff", bd=2, relief=tk.GROOVE)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        self.text_area = scrolledtext.ScrolledText(
            text_frame,
            wrap=tk.WORD,
            width=70,
            height=18,
            state='disabled',
            font=self.custom_font,
            bg="#ffffff",
            fg="#333333",
            padx=15,
            pady=15
        )
        self.text_area.pack(fill=tk.BOTH, expand=True)
        
        # 添加初始提示
        self.display_text("欢迎使用AI面试智能官！\n请按以下步骤操作：\n1. 上传简历\n2. 上传JD职位描述\n3. 选择面试赛道\n4. 开始面试\n\n面试过程中系统会智能分析您的回答，请放心作答。")

        # 创建录音控制区域
        control_frame = tk.Frame(main_frame, bg="#f0f0f0")
        control_frame.pack(fill=tk.X, pady=10)
        
        # 添加录音状态标签
        self.status_label = tk.Label(
            control_frame,
            text="准备就绪",
            font=self.small_font,
            bg="#f0f0f0",
            fg="#7f8c8d"
        )
        self.status_label.pack(side=tk.LEFT, padx=10)
        
        # 添加进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            control_frame,
            variable=self.progress_var,
            maximum=100,
            length=300,
            mode='determinate'
        )
        self.progress_bar.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10)
        
        # 添加录音时间标签
        self.time_label = tk.Label(
            control_frame,
            text="0.0s",
            font=self.small_font,
            bg="#f0f0f0",
            fg="#7f8c8d",
            width=6
        )
        self.time_label.pack(side=tk.LEFT, padx=10)
        
        # 添加操作提示
        help_frame = tk.Frame(main_frame, bg="#f0f0f0")
        help_frame.pack(fill=tk.X, pady=5)
        
        self.help_label = tk.Label(
            help_frame,
            text="按住空格键开始录音，松开结束录音",
            font=("Helvetica", 12, "italic"),
            bg="#f0f0f0",
            fg="#3498db"
        )
        self.help_label.pack()

        # 绑定按键事件
        self.root.bind("<KeyPress-space>", self.start_recording)
        self.root.bind("<KeyRelease-space>", self.stop_recording)
        self.root.focus_set()

        # 创建队列，用于线程间通信
        self.message_queue = queue.Queue()
        self.input_queue = queue.Queue()

        # 启动队列监听器
        self.start_queue_listener()
        
        # 启动模型响应线程
        self.model_thread = threading.Thread(target=self.process_model_responses, daemon=True)
        self.model_thread.start()
        
        # 检查初始状态
        self.check_interview_ready()
        
        # 状态变量
        self.recording_start_time = 0
        self.progress_active = False
        self.is_processing = False
        self.interview_active = False
        self.resume_data = None
        self.parser = ResumeParser()
        
        # 对话历史
        self.conversation_history = []
        self.initial_prompt_set = False
        self.question_count = 0  # 问题计数器
        self.first_question_asked = False  # 标记是否已问过第一个问题
        self.key_topics = []  # 存储关键话题用于动态提示调整

    def select_track(self):
        """显示赛道选择对话框"""
        track_window = tk.Toplevel(self.root)
        track_window.title("选择面试赛道")
        track_window.geometry("400x300")
        track_window.configure(bg="#f0f0f0")
        track_window.transient(self.root)
        track_window.grab_set()
        
        # 标题
        title_label = tk.Label(
            track_window,
            text="请选择面试赛道",
            font=("Helvetica", 16, "bold"),
            bg="#f0f0f0",
            fg="#2c3e50"
        )
        title_label.pack(pady=20)
        
        # 赛道按钮
        tracks = ["后端", "前端", "算法", "测试", "产品", "运营"]
        colors = ["#e74c3c", "#3498db", "#2ecc71", "#f39c12", "#9b59b6", "#1abc9c"]
        
        for i, track in enumerate(tracks):
            btn = tk.Button(
                track_window,
                text=f"🎯 {track}",
                command=lambda t=track: self.confirm_track_selection(t, track_window),
                font=("Helvetica", 12, "bold"),
                bg=colors[i % len(colors)],
                fg="white",
                padx=20,
                pady=10,
                relief=tk.FLAT,
                width=15
            )
            btn.pack(pady=5)
    
    def confirm_track_selection(self, track, window):
        """确认赛道选择"""
        self.selected_track = track
        self.question_bank_manager.set_track(track)
        
        # 更新界面显示
        track_summary = self.question_bank_manager.get_track_summary()
        self.track_info_label.config(
            text=f"已选择: {track} 赛道\n{track_summary}",
            fg="#2ecc71"
        )
        
        # 更新赛道选择状态
        self.check_interview_ready()
        
        self.display_text(f"已选择 {track} 赛道！题库已加载完成。")
        
        window.destroy()
    
    def upload_jd(self):
        """上传并解析JD"""
        # 创建JD输入对话框
        jd_window = tk.Toplevel(self.root)
        jd_window.title("输入职位描述(JD)")
        jd_window.geometry("600x400")
        jd_window.configure(bg="#f0f0f0")
        jd_window.transient(self.root)
        jd_window.grab_set()
        
        # 标题
        title_label = tk.Label(
            jd_window,
            text="请输入职位描述(JD)信息",
            font=("Helvetica", 14, "bold"),
            bg="#f0f0f0",
            fg="#2c3e50"
        )
        title_label.pack(pady=10)
        
        # 文本输入框
        jd_text = scrolledtext.ScrolledText(
            jd_window,
            wrap=tk.WORD,
            width=70,
            height=20,
            font=("Helvetica", 10)
        )
        jd_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # 按钮框架
        btn_frame = tk.Frame(jd_window, bg="#f0f0f0")
        btn_frame.pack(pady=10)
        
        def confirm_jd():
            jd_content = jd_text.get("1.0", tk.END).strip()
            if jd_content:
                self.jd_data = self.jd_analyzer.parse_jd(jd_content)
                jd_summary = self.jd_analyzer.get_jd_summary()
                self.display_text(f"JD解析完成！\n{jd_summary}")
                
                # 检查是否可以开始面试
                self.check_interview_ready()
                    
                jd_window.destroy()
            else:
                tk.messagebox.showerror("错误", "请输入JD内容")
        
        # 确认按钮
        confirm_btn = tk.Button(
            btn_frame,
            text="确认",
            command=confirm_jd,
            font=("Helvetica", 12),
            bg="#2ecc71",
            fg="white",
            padx=20,
            pady=5
        )
        confirm_btn.pack(side=tk.LEFT, padx=10)
        
        # 取消按钮
        cancel_btn = tk.Button(
            btn_frame,
            text="取消",
            command=jd_window.destroy,
            font=("Helvetica", 12),
            bg="#e74c3c",
            fg="white",
            padx=20,
            pady=5
        )
        cancel_btn.pack(side=tk.LEFT, padx=10)

    def upload_resume(self):
        """上传并解析简历"""
        file_path = filedialog.askopenfilename(
            title="选择简历文件",
            filetypes=[("PDF文件", "*.pdf"), ("Word文件", "*.docx"), ("所有文件", "*.*")]
        )
        
        if file_path:
            self.display_text(f"正在解析简历: {os.path.basename(file_path)}...")
            self.resume_data = self.parser.parse_resume(file_path)
            
            if isinstance(self.resume_data, dict):
                # 更新候选人信息面板
                info_text = f"姓名: {self.resume_data['name']}\n"
                info_text += f"联系方式: {self.resume_data['contact']}\n"
                info_text += f"教育背景: {len(self.resume_data['education'])}项\n"
                info_text += f"工作经验: {len(self.resume_data['experience'])}项\n"
                info_text += f"技能: {len(self.resume_data['skills'])}项"
                
                self.info_label.config(text=info_text)
                self.display_text("简历解析完成！")
                
                # 更新简历状态
                self.resume_data = self.resume_data
                
                # 创建动态提示调整器
                self.dynamic_prompt_adjuster = DynamicPromptAdjuster(self.resume_data)
                
                # 从简历中提取关键实体
                entities = self.parser.extract_entities(file_path)  # 修复这里：使用文件路径提取实体
                if entities:
                    self.dynamic_prompt_adjuster.key_entities = entities
                    self.display_text(f"已提取关键实体: {', '.join(entities[:5])}...")
                else:
                    self.display_text("未提取到关键实体。")
            else:
                self.display_text(f"简历解析失败: {self.resume_data}")
    
    def extract_key_topics(self):
        """从简历中提取关键话题用于动态提示调整"""
        if not self.resume_data:
            return
            
        # 提取关键技能
        self.key_topics = []
        if self.resume_data.get("skills"):
            self.key_topics.extend([skill for skill in self.resume_data["skills"] if len(skill) > 3])
            
        # 提取项目经验中的关键技术
        if self.resume_data.get("projects"):
            for project in self.resume_data["projects"]:
                if "(" in project and ")" in project:
                    techs = project.split("(")[1].split(")")[0].split(",")
                    self.key_topics.extend([tech.strip() for tech in techs if len(tech.strip()) > 3])
                    
        # 去重
        self.key_topics = list(set(self.key_topics))
        self.display_text(f"已提取关键话题: {', '.join(self.key_topics[:5])}...")

    def start_interview(self):
        """开始面试"""
        if not self.resume_data:
            self.display_text("请先上传简历！")
            return
        
        if not self.selected_track:
            self.display_text("请先选择面试赛道！")
            return
            
        if not self.jd_data:
            self.display_text("请先上传JD！")
            return
        
        self.interview_active = True
        self.question_count = 0  # 重置问题计数器
        self.first_question_asked = False  # 重置第一问题标记
        
        # 重置评分和难度管理器
        self.score_manager = ScoreAndDifficultyManager()
        
        # 重置三阶段管理器
        self.stage_manager = ThreeStageInterviewManager()
        
        # 重置复盘管理器
        self.review_manager = InterviewReviewManager()
        
        # 更新状态显示
        self.update_status_display()
        current_stage = self.stage_manager.get_current_stage()
        self.stage_info_label.config(text=f"当前阶段: {current_stage} (第1题)")
        
        self.display_text("面试已开始！请准备回答面试官的问题。")
        self.display_text(f"📍 当前阶段: {current_stage}")
        self.end_interview_btn.config(state=tk.NORMAL)
        self.start_interview_btn.config(state=tk.DISABLED)
        
        # 禁用复盘和导出按钮
        self.review_btn.config(state=tk.DISABLED)
        self.export_pdf_btn.config(state=tk.DISABLED)
        
        # 重置对话历史
        self.conversation_history = []
        self.initial_prompt_set = False
        
        # 发送初始面试问题请求
        self.input_queue.put("start_interview")

    def end_interview(self):
        """结束面试"""
        self.interview_active = False
        self.display_text("面试已结束！感谢参与。")
        
        # 显示评分总结
        score_summary = self.score_manager.get_score_summary()
        difficulty_progression = self.score_manager.get_difficulty_progression()
        
        self.display_text("\n" + "="*50)
        self.display_text("📊 面试评分总结")
        self.display_text("="*50)
        self.display_text(score_summary)
        self.display_text(difficulty_progression)
        self.display_text("="*50)
        
        self.end_interview_btn.config(state=tk.DISABLED)
        self.start_interview_btn.config(state=tk.NORMAL)
        
        # 发送评估请求（包含评分信息）
        self.input_queue.put("end_interview")
        
        # 启用复盘和导出按钮
        self.review_btn.config(state=tk.NORMAL)
        self.export_pdf_btn.config(state=tk.NORMAL)
    
    def show_interview_review(self):
        """显示面试复盘"""
        if not self.review_manager.interview_records:
            tk.messagebox.showinfo("提示", "暂无面试记录")
            return
        
        # 获取完整评估
        if hasattr(self, 'full_evaluation'):
            assessment = self.full_evaluation
        else:
            # 如果没有完整评估，生成一个
            assessment = self.review_manager.generate_overall_assessment(
                self.stage_manager, self.score_manager
            )
        
        # 创建复盘窗口
        review_window = tk.Toplevel(self.root)
        review_window.title("面试复盘分析")
        review_window.geometry("800x600")
        review_window.configure(bg="#f0f0f0")
        
        # 创建滚动文本框
        review_text = scrolledtext.ScrolledText(
            review_window,
            wrap=tk.WORD,
            width=90,
            height=35,
            font=("Helvetica", 10),
            bg="#ffffff",
            fg="#333333"
        )
        review_text.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # 显示复盘内容
        review_content = f"""
🎯 面试复盘报告
{'='*60}

📅 面试日期: {assessment['interview_date']}
📊 总题数: {assessment['total_questions']}题
🎯 综合得分: {assessment['overall_score']:.2f}/1.00

{'='*60}
📈 各阶段表现分析
{'='*60}
"""
        
        for stage, performance in assessment['stage_performance'].items():
            review_content += f"""
🔸 {stage}:
   - 题目数量: {performance['questions']}题
   - 平均得分: {performance['avg_score']:.2f}/1.00
   - 表现评价: {'优秀' if performance['avg_score'] >= 0.75 else '良好' if performance['avg_score'] >= 0.6 else '一般' if performance['avg_score'] >= 0.5 else '需改进'}
"""
        
        review_content += f"""
{'='*60}
✅ 优势表现
{'='*60}
"""
        for strength in assessment['strengths']:
            review_content += f"• {strength}\n"
        
        if not assessment['strengths']:
            review_content += "建议在各个方面继续努力\n"
        
        review_content += f"""
{'='*60}
⚠️ 待改进方面
{'='*60}
"""
        for weakness in assessment['weaknesses']:
            review_content += f"• {weakness}\n"
        
        if not assessment['weaknesses']:
            review_content += "整体表现良好\n"
        
        review_content += f"""
{'='*60}
💡 改进建议
{'='*60}
"""
        for suggestion in assessment['improvement_suggestions']:
            review_content += f"• {suggestion}\n"
        
        review_content += f"""
{'='*60}
📝 详细问答记录
{'='*60}
"""
        
        for i, record in enumerate(self.review_manager.interview_records):
            review_content += f"""
第{i+1}题 [{record['stage']}] - {record['timestamp']}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❓ 问题: {record['question']}

💬 回答: {record['answer']}

📊 得分: {record['score']:.2f}/1.00

💭 反馈: {record['feedback']}

"""
        
        review_text.insert(tk.END, review_content)
        review_text.config(state='disabled')
    
    def export_interview_pdf(self):
        """导出面试记录为PDF"""
        if not self.review_manager.interview_records:
            tk.messagebox.showinfo("提示", "暂无面试记录")
            return
        
        # 选择保存位置
        filename = filedialog.asksaveasfilename(
            title="保存面试报告",
            defaultextension=".pdf",
            filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")]
        )
        
        if filename:
            # 生成综合评估
            self.review_manager.generate_overall_assessment(
                self.stage_manager, self.score_manager
            )
            
            # 导出PDF
            candidate_name = self.resume_data.get('name', '候选人') if self.resume_data else '候选人'
            track_name = self.selected_track or '未知赛道'
            
            success = self.review_manager.export_to_pdf(filename, candidate_name, track_name)
            
            if success:
                tk.messagebox.showinfo("成功", f"面试报告已导出至:\n{filename}")
            else:
                tk.messagebox.showerror("失败", "PDF导出失败，请检查权限和路径")

    def start_recording(self, event):
        """开始录音"""
        if not self.interview_active:
            self.display_text("请先开始面试！")
            return
            
        if not self.is_processing and not self.progress_active:
            self.solution.recorder.start_recording()
            self.recording_start_time = time.time()
            self.progress_active = True
            self.update_progress()
            self.status_label.config(text="录音中...", fg="#e74c3c")

    def update_progress(self):
        if not self.progress_active:
            return
            
        elapsed = time.time() - self.recording_start_time
        progress = min(100, math.log(1 + elapsed * 10) * 30)
        self.progress_var.set(progress)
        self.time_label.config(text=f"{elapsed:.1f}s")
        
        if self.progress_active:
            self.root.after(50, self.update_progress)

    def stop_recording(self, event):
        """停止录音并处理录音内容"""
        if not self.progress_active or self.is_processing or not self.interview_active:
            return
            
        self.progress_active = False
        self.is_processing = True
        self.status_label.config(text="处理中...", fg="#f39c12")
        user_input = self.solution.recorder.stop_recording()
        
        if user_input:
            self.message_queue.put(f"候选人: {user_input}")
            
            # 保存最后回答用于评分
            self.last_answer = user_input
            
            # 将候选人的回答添加到对话历史
            self.conversation_history.append({"role": "user", "content": user_input})
            
            # 分析回答并提取关键点
            self.analyze_response(user_input)
            
            # 发送处理请求
            self.input_queue.put("candidate_response")
        
        self.root.after(100, self.reset_progress)

    def analyze_response(self, response):
        """分析候选人回答，更新对话上下文，评分并调整难度"""
        # 更新动态提示调整器
        if self.dynamic_prompt_adjuster and hasattr(self, 'last_model_output'):
            self.dynamic_prompt_adjuster.update_conversation_context(
                response, self.last_model_output
            )
        
        # 添加到对话上下文
        self.conversation_context.append(f"候选人: {response}")
        
        # 限制上下文长度
        if len(self.conversation_context) > 6:
            self.conversation_context = self.conversation_context[-6:]
        
        # 使用评分系统对回答进行评分（只有在面试已经开始且有问题的情况下）
        if (hasattr(self, 'last_question') and self.last_question and 
            hasattr(self, 'question_count') and self.question_count > 0 and
            self.interview_active):
            # 在后台静默进行评分，不显示给用户
            
            def score_response():
                try:
                    score = self.score_manager.evaluate_response(response, self.last_question)
                    
                    # 根据评分调整难度
                    new_difficulty = self.score_manager.adjust_difficulty(score)
                    
                    # 在主线程中更新后台状态
                    self.root.after(0, lambda: self._update_after_scoring_silent(score, new_difficulty))
                
                except Exception as e:
                    print(f"评分过程出错: {e}")
                    # 评分失败也不显示给用户，只在控制台记录
            
            threading.Thread(target=score_response, daemon=True).start()
    
    def _update_after_scoring_silent(self, score, new_difficulty):
        """静默评分完成后更新后台状态，不显示给用户"""
        # 生成详细反馈
        if hasattr(self, 'last_question') and hasattr(self, 'last_answer'):
            feedback = self.review_manager.generate_detailed_feedback(
                self.last_answer, self.last_question, score
            )
        else:
            feedback = f"得分: {score:.2f}/1.0"
        
        # 添加到复盘记录
        current_stage = self.stage_manager.get_current_stage()
        self.review_manager.add_qa_record(
            self.last_question, self.last_answer, score, feedback, current_stage
        )
        
        # 添加评分到当前阶段
        self.stage_manager.add_score_to_stage(score)
        
        # 检查是否需要进入下一阶段
        if self.stage_manager.should_advance_stage():
            if self.stage_manager.advance_to_next_stage():
                new_stage = self.stage_manager.get_current_stage()
                # 静默切换阶段，不显示给用户
                self.stage_info_label.config(
                    text=f"当前阶段: {new_stage} (第{self.stage_manager.current_stage_question_count + 1}题)"
                )
            else:
                # 所有阶段完成，不显示给用户
                self.stage_info_label.config(text="当前阶段: 面试完成")
        else:
            current_stage = self.stage_manager.get_current_stage()
            self.stage_info_label.config(
                text=f"当前阶段: {current_stage} (第{self.stage_manager.current_stage_question_count + 1}题)"
            )
        
        # 动态调整后续阶段问题数量
        self.stage_manager.adjust_stage_questions(self.score_manager.score_history)
        
        # 更新状态显示面板（仅更新技术指标，不显示评分）
        self.update_status_display_silent()
    
    def _update_after_scoring(self, score, new_difficulty):
        """评分完成后更新UI和状态"""
        # 生成详细反馈
        if hasattr(self, 'last_question') and hasattr(self, 'last_answer'):
            feedback = self.review_manager.generate_detailed_feedback(
                self.last_answer, self.last_question, score
            )
        else:
            feedback = f"得分: {score:.2f}/1.0"
        
        # 添加到复盘记录
        current_stage = self.stage_manager.get_current_stage()
        self.review_manager.add_qa_record(
            self.last_question, self.last_answer, score, feedback, current_stage
        )
        
        # 显示评分结果
        score_text = f"本题评分: {score:.2f}"
        if score > 0.8:
            score_text += " (优秀)"
        elif score > 0.6:
            score_text += " (良好)"
        elif score > 0.4:
            score_text += " (及格)"
        else:
            score_text += " (需改进)"
        
        self.display_text(score_text)
        self.display_text(f"反馈: {feedback}")
        
        # 添加评分到当前阶段
        self.stage_manager.add_score_to_stage(score)
        
        # 检查是否需要进入下一阶段
        if self.stage_manager.should_advance_stage():
            if self.stage_manager.advance_to_next_stage():
                new_stage = self.stage_manager.get_current_stage()
                self.display_text(f"📍 进入下一阶段: {new_stage}")
                self.stage_info_label.config(
                    text=f"当前阶段: {new_stage} (第{self.stage_manager.current_stage_question_count + 1}题)"
                )
            else:
                self.display_text("📍 所有阶段已完成，可以结束面试")
                self.stage_info_label.config(text="当前阶段: 面试完成")
        else:
            current_stage = self.stage_manager.get_current_stage()
            self.stage_info_label.config(
                text=f"当前阶段: {current_stage} (第{self.stage_manager.current_stage_question_count + 1}题)"
            )
        
        # 动态调整后续阶段问题数量
        self.stage_manager.adjust_stage_questions(self.score_manager.score_history)
        
        # 如果难度发生变化，显示难度调整信息
        if new_difficulty != self.score_manager.difficulty_history[-1]["previous_difficulty"]:
            self.display_text(f"难度调整: {self.score_manager.difficulty_history[-1]['previous_difficulty']} → {new_difficulty}")
        
        # 更新状态显示
        self.update_status_display()
    
    def update_status_display_silent(self):
        """静默更新面试状态显示面板（不显示评分信息）"""
        # 更新当前难度（不显示给用户，保持原有显示）
        # self.difficulty_label.config(text=f"当前难度: {self.score_manager.current_difficulty}")
        
        # 更新问题计数
        if hasattr(self, 'question_count'):
            self.question_count_label.config(text=f"问题数: {self.question_count}")
        
        # 显示通用的进度信息，不显示具体评分
        self.latest_score_label.config(text="面试进行中...")
        self.avg_score_label.config(text="实时分析中...")
    
    def update_status_display(self):
        """更新面试状态显示面板"""
        # 更新当前难度
        self.difficulty_label.config(text=f"当前难度: {self.score_manager.current_difficulty}")
        
        # 更新最新评分
        if self.score_manager.score_history:
            latest_score = self.score_manager.score_history[-1]
            self.latest_score_label.config(text=f"最新评分: {latest_score:.2f}")
            
            # 更新平均分
            avg_score = sum(self.score_manager.score_history) / len(self.score_manager.score_history)
            self.avg_score_label.config(text=f"平均分: {avg_score:.2f}")
        else:
            self.latest_score_label.config(text="最新评分: --")
            self.avg_score_label.config(text="平均分: --")
        
        # 更新问题计数
        self.question_count_label.config(text=f"问题数: {len(self.score_manager.score_history)}")

    def reset_progress(self):
        self.progress_var.set(0)
        self.time_label.config(text="0.0s")

    def extract_question(self, model_output):
        """从模型输出中提取问题部分，过滤掉思考过程"""
        # 移除<think>标签之间的内容
        while "<think>" in model_output and "</think>" in model_output:
            think_start = model_output.find("<think>")
            think_end = model_output.find("</think>") + len("</think>")
            model_output = model_output[:think_start] + model_output[think_end:]
        
        # 寻找 ">" 符号
        arrow_index = model_output.find(">")
        if arrow_index != -1:
            # 提取 ">" 之后的内容
            question_text = model_output[arrow_index + 1:].strip()
            # 如果问题后面还有内容，只保留到第一个句号或问号
            end_marks = ['.', '?', '！', '。', '？']
            end_indices = [question_text.find(mark) for mark in end_marks if mark in question_text]
            if end_indices:
                valid_indices = [i for i in end_indices if i != -1]
                if valid_indices:
                    first_end = min(valid_indices) + 1
                    question_text = question_text[:first_end].strip()
            return question_text
        
            # 如果没有找到 ">"，尝试找到最后一个问号之后的内容
            last_question = model_output.rfind("?")
            if last_question != -1:
            return model_output[:last_question + 1].strip()
            
            # 如果还是没有找到，返回整个输出
        return model_output.strip()

    def build_dynamic_prompt(self):
        """根据论文结构构建动态提示 [I; H; K; E]，集成三阶段面试和难度调整"""
        # 获取当前阶段提示
        stage_prompt = self.stage_manager.get_stage_prompt(self.jd_data, self.resume_data)
        
        # 获取难度相关提示
        difficulty_prompt = self.score_manager.get_difficulty_prompt()
        
        # 获取题库参考问题
        reference_questions = []
        if self.selected_track:
            ref_questions = self.question_bank_manager.get_reference_questions(
                difficulty=self.score_manager.current_difficulty,
                num_questions=3
            )
            reference_questions = [q["question"] for q in ref_questions]
        
        # I: 任务指令（集成三阶段面试、难度调整和题库）
        instruction = (
            f"你是一个专业的{self.selected_track}面试官，具备三阶段面试和动态难度调整能力。"
            f"当前正在进行第{self.stage_manager.current_stage + 1}阶段面试。\n\n"
            f"{stage_prompt}\n\n"
            f"{difficulty_prompt}\n\n"
        )
        
        # 添加JD相关信息
        if self.jd_data and self.jd_data.get('keywords'):
            instruction += f"### JD关键技术要求:\n{', '.join(self.jd_data['keywords'])}\n\n"
        
        # 添加题库参考
        if reference_questions:
            instruction += "### 参考题库示例:\n"
            for i, q in enumerate(reference_questions, 1):
                instruction += f"{i}. {q}\n"
            instruction += "\n请参考以上题库内容和当前阶段要求，生成相应的问题。\n\n"
        
        instruction += (
            "你必须严格遵守以下规则：\n"
            "1. 直接输出问题，不要输出思考过程\n"
            "2. 在问题前添加'>'符号作为前缀\n"
            "3. 只输出问题内容，不要添加任何前缀（如'面试官：'）\n"
            "4. 每次只提一个问题\n"
            "5. 问题应该简洁明了，不超过2句话\n"
            "6. 问题难度必须与当前设定的难度级别匹配\n"
            f"7. 问题必须符合当前{self.stage_manager.get_current_stage()}阶段要求\n"
            f"8. 问题必须符合{self.selected_track}赛道的专业要求\n"
            "9. 面试结束时只评价候选人表现，不要评价问题质量\n"
        )
        
        # H: 历史细节
        history_details = "### 对话历史摘要:\n"
        if self.conversation_context:
            history_details += "\n".join(self.conversation_context[-3:])
        else:
            history_details += "暂无历史对话"
        
        # K: 证据细节（知识）
        evidence_details = "### 证据细节:\n"
        if self.dynamic_prompt_adjuster:
            # 应用三元组过滤器
            triplets = self.dynamic_prompt_adjuster.triplet_filter()
            if triplets:
                evidence_details += "知识三元组:\n"
                for head, rel, tail in triplets:
                    evidence_details += f"- {head} -- {rel} --> {tail}\n"
            
            # 添加上下文实体
            if self.dynamic_prompt_adjuster.historical_entities:
                evidence_details += "\n历史实体:\n"
                evidence_details += ", ".join(set(self.dynamic_prompt_adjuster.historical_entities[-5:]))
            
            # 添加当前实体
            if self.dynamic_prompt_adjuster.current_entities:
                evidence_details += "\n当前相关实体:\n"
                evidence_details += ", ".join(self.dynamic_prompt_adjuster.current_entities)
        
        # E: 相关演示
        demonstration = "### 相关演示:\n"
        if self.dynamic_prompt_adjuster:
            demo = self.dynamic_prompt_adjuster.demo_selector()
            demonstration += demo
        
        # 组合所有部分
        prompt = f"""
        {instruction}
        
        {history_details}
        
        {evidence_details}
        
        {demonstration}
        """
        
        return prompt

    def process_model_responses(self):
        """处理模型响应的线程"""
        while True:
            action = self.input_queue.get()
            
            if not self.interview_active and action != "start_interview":
                continue
            
            try:
                # 构建动态提示
                dynamic_prompt = self.build_dynamic_prompt()
                
                # 设置系统提示
                if not self.initial_prompt_set:
                    self.conversation_history = [
                        {"role": "system", "content": dynamic_prompt}
                    ]
                    self.initial_prompt_set = True
                else:
                    # 更新系统提示
                    self.conversation_history[0] = {"role": "system", "content": dynamic_prompt}
                
                # 处理不同操作
                if action == "start_interview":
                    # 添加第一个问题请求
                    self.conversation_history.append({
                        "role": "user", 
                        "content": "请基于候选人的简历提出第一个面试问题"
                    })
                    self.question_count = 0
                    self.first_question_asked = False
                elif action == "end_interview":
                    # 添加评估请求（包含评分信息）
                    score_info = self.score_manager.get_score_summary()
                    difficulty_info = self.score_manager.get_difficulty_progression()
                    
                    evaluation_content = f"""面试结束，请根据整个面试过程给出候选人综合评估。
请从以下几个方面进行评价：

1. 技术掌握程度
2. 项目经验深度
3. 解决问题能力
4. 沟通表达能力
5. 学习成长潜力

注意：
- 只评价候选人的表现，不要评价问题的质量
- 不要显示具体分数，用定性描述替代
- 重点关注候选人的优势和可提升空间
- 给出明确的是否通过面试的结论"""

                    self.conversation_history.append({
                        "role": "user", 
                        "content": evaluation_content
                    })
                elif action == "candidate_response":
                    # 候选人的回答已经在对话历史中，不需要额外处理
                    pass
                
                # 调用模型生成回复
                output = ollama.chat(
                    model="Jerrypoi/deepseek-r1-with-tool-calls:latest",
                    messages=self.conversation_history
                )
                model_output = output['message']['content']
                self.last_model_output = model_output  # 保存最后输出
                
                # 将模型的完整回复添加到对话历史
                self.conversation_history.append({"role": "assistant", "content": model_output})
                
                # 如果是结束面试的评估，只显示简短总结
                if action == "end_interview":
                    # 保存完整评估结果用于PDF导出
                    self.full_evaluation = model_output
                    
                    # 只显示简短的结论
                    self.message_queue.put("面试已结束！请点击"查看复盘"按钮查看详细评估，或导出PDF报告查看完整分析。")
                    self.solution.use_pyttsx3("面试评估已完成")
                    continue
                
                # 提取问题部分（">"之后的内容）
                question_text = self.extract_question(model_output)
                
                # 保存最后一个问题用于评分
                self.last_question = question_text
                
                # 显示问题 - 只显示">"之后的内容
                self.message_queue.put(f"> {question_text}")
                
                # 语音播报问题 - 只播报">"之后的内容
                self.solution.use_pyttsx3(question_text)
                
                # 更新状态
                self.status_label.config(text="回答中...", fg="#9b59b6")
                
                # 增加问题计数
                self.question_count += 1
                
                # 如果是第一次提问后，更新系统提示为基于回答的追问
                if not self.first_question_asked and self.question_count >= 1:
                    self.first_question_asked = True
                    self.display_text("面试进入深入追问阶段...")
                
            except Exception as e:
                error_msg = f"错误: {str(e)}"
                self.message_queue.put(error_msg)
                self.solution.use_pyttsx3("抱歉，处理您的请求时出错了")
            
            finally:
                self.is_processing = False
                if self.interview_active:
                    self.status_label.config(text="准备提问...", fg="#2ecc71")
                else:
                    self.status_label.config(text="准备就绪", fg="#2ecc71")

    def start_queue_listener(self):
        """监听队列，将消息显示在文本框中"""
        def check_queue():
            try:
                while True:
                    message = self.message_queue.get_nowait()
                    self.display_text(message)
            except queue.Empty:
                pass
            self.root.after(100, check_queue)

        check_queue()

    def check_interview_ready(self):
        """检查是否满足开始面试的条件"""
        remaining = []
        if not self.resume_data:
            remaining.append("上传简历")
        if not self.selected_track:
            remaining.append("选择赛道")
        if not self.jd_data:
            remaining.append("上传JD")
            
        if not remaining:  # 所有条件都满足
            self.start_interview_btn.config(state=tk.NORMAL)
            self.display_text("✅ 已完成所有准备工作，可以开始面试！")
        else:
            self.start_interview_btn.config(state=tk.DISABLED)
            if len(remaining) > 1:
                self.display_text(f"⏳ 请完成以下步骤：{', '.join(remaining)}")
            else:
                self.display_text(f"⏳ 请{remaining[0]}后开始面试")

    def display_text(self, text):
        """将文本显示在文本框中"""
        self.text_area.config(state='normal')
        
        if text.startswith("候选人:"):
            self.text_area.tag_configure("candidate", foreground="#2980b9", font=(self.default_family, 14, "bold"))
            self.text_area.insert(tk.END, text + "\n\n", "candidate")
        elif text.startswith(">"):
            self.text_area.tag_configure("interviewer", foreground="#27ae60", font=(self.default_family, 14))
            self.text_area.insert(tk.END, text + "\n\n", "interviewer")
        elif text.startswith("错误:") or text.startswith("面试评估:"):
            self.text_area.tag_configure("error", foreground="#e74c3c", font=(self.default_family, 14))
            self.text_area.insert(tk.END, text + "\n", "error")
        else:
            self.text_area.insert(tk.END, text + "\n")
        
        self.text_area.config(state='disabled')
        self.text_area.yview(tk.END)
# 主逻辑模块
class Solution:
    def __init__(self):
        model = Model("vosk-model-cn-0.15")
        SetLogLevel(-1)
        self.recorder = VoiceRecorder(model)

    def use_pyttsx3(self, word):
        """语音播报"""
        threading.Thread(target=self._speak, args=(word,), daemon=True).start()
    
    def _speak(self, word):
        try:
            engine = pyttsx3.init()
            engine.setProperty('rate', 200)
            engine.setProperty('volume', 1.0)
            voices = engine.getProperty('voices')
            if voices:
                engine.setProperty('voice', voices[0].id)
            engine.say(word)
            engine.runAndWait()
            engine.stop()
        except Exception as e:
            print(f"语音播报出错: {e}")

# 主程序
if __name__ == "__main__":
    # 创建解决方案实例
    solution = Solution()
    
    # 创建 Tkinter 窗口
    root = tk.Tk()
    app = InteractiveTextApp(root, solution)
    
    # 运行 Tkinter 主循环
    root.mainloop()
