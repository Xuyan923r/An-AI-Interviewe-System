# -*- coding: utf-8 -*-

import re
import fitz
import docx


class ResumeParser:
    def __init__(self):
        self.resume_data = {
            "name": "",
            "contact": "",
            "education": [],
            "experience": [],
            "skills": [],
            "projects": [],
            "summary": ""
        }

    def parse_resume(self, file_path):
        if file_path.lower().endswith('.pdf'):
            return self.parse_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.parse_docx(file_path)
        else:
            return "不支持的文件格式，请上传PDF或DOCX文件"

    def parse_pdf(self, file_path):
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return self._extract_info(text)
        except Exception as e:
            return f"解析PDF失败: {str(e)}"

    def parse_docx(self, file_path):
        try:
            d = docx.Document(file_path)
            text = "\n".join([para.text for para in d.paragraphs])
            return self._extract_info(text)
        except Exception as e:
            return f"解析DOCX失败: {str(e)}"

    def _extract_info(self, text):
        name_match = re.search(r"^(.*?)\n", text)
        if name_match:
            self.resume_data["name"] = name_match.group(1).strip()

        phone_match = re.search(r"(\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4})", text)
        email_match = re.search(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})", text)
        self.resume_data["contact"] = f"电话: {phone_match.group(0) if phone_match else '未找到'} | 邮箱: {email_match.group(0) if email_match else '未找到'}"

        edu_section = re.search(r"(教育背景|教育经历|学历|Education)(.*?)(?=(工作经历|项目经历|技能|$))", text, re.DOTALL | re.IGNORECASE)
        if edu_section:
            edu_text = edu_section.group(2)
            self.resume_data["education"] = [line.strip() for line in edu_text.split('\n') if line.strip()]

        exp_section = re.search(r"(工作经历|工作经验|工作|Experience)(.*?)(?=(项目经历|技能|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if exp_section:
            exp_text = exp_section.group(2)
            self.resume_data["experience"] = [line.strip() for line in exp_text.split('\n') if line.strip()]

        skills_section = re.search(r"(技能|专业技能|技术能力|Skills)(.*?)(?=(项目经历|工作经历|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if skills_section:
            skills_text = skills_section.group(2)
            self.resume_data["skills"] = [line.strip() for line in skills_text.split('\n') if line.strip()]

        projects_section = re.search(r"(项目经历|项目经验|项目|Projects)(.*?)(?=(技能|工作经历|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if projects_section:
            projects_text = projects_section.group(2)
            self.resume_data["projects"] = [line.strip() for line in projects_text.split('\n') if line.strip()]

        self.resume_data["summary"] = f"候选人: {self.resume_data['name']}，教育背景: {len(self.resume_data['education'])}项，工作经验: {len(self.resume_data['experience'])}项，技能: {len(self.resume_data['skills'])}项"
        return self.resume_data

    def extract_entities(self, file_path):
        text = ""
        if file_path.lower().endswith('.pdf'):
            try:
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text += page.get_text()
            except Exception:
                return []
        elif file_path.lower().endswith('.docx'):
            try:
                d = docx.Document(file_path)
                text = "\n".join([para.text for para in d.paragraphs])
            except Exception:
                return []
        else:
            return []

        entities = []
        skills_section = re.search(r"(技能|专业技能|技术能力|Skills)(.*?)(?=(项目经历|工作经历|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if skills_section:
            skills_text = skills_section.group(2)
            skills = [line.strip() for line in skills_text.split('\n') if line.strip()]
            entities.extend(skills)

        projects_section = re.search(r"(项目经历|项目经验|项目|Projects)(.*?)(?=(技能|工作经历|教育背景|$))", text, re.DOTALL | re.IGNORECASE)
        if projects_section:
            projects_text = projects_section.group(2)
            projects = [line.strip() for line in projects_text.split('\n') if line.strip()]
            for project in projects:
                if "(" in project and ")" in project:
                    techs = project.split("(")[1].split(")")[0].split(",")
                    entities.extend([tech.strip() for tech in techs])
        return list(set(entities))



